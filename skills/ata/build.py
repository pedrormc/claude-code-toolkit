"""
build.py — Gera uma Ata de Reunião Singular em .docx a partir de um JSON.

Uso:
    python build.py content.json output.docx

Reutiliza o template.docx oficial da Singular (herdado de ~/.claude/skills/pop/template.docx)
que já tem fonts Urbanist, header com logo e footer paginado. Limpa o corpo do
documento e reconstrói com a estrutura de ata (metadados, participantes, tópicos,
encaminhamentos, próxima reunião).

Estrutura esperada do content.json — ver SKILL.md.
"""
import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

LOGO_PATH = Path(__file__).parent / "assets" / "logo-singular.png"

URBANIST = "Urbanist"
PRETO = RGBColor(0x1C, 0x1C, 0x1C)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA = RGBColor(0x55, 0x55, 0x55)
CINZA_CLARO = RGBColor(0xD0, 0xD0, 0xD0)


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="d0d0d0", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def set_font(run, name=URBANIST, size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def clear_document_body(doc):
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def add_paragraph(doc, text, size=11, bold=False, align=None, color=None, space_after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color or PRETO)
    return p


def add_logo(doc):
    """Insere a logo oficial da Singular centralizada no topo do documento."""
    if not LOGO_PATH.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(1.8))


def add_title_block(doc, titulo_curto, titulo):
    add_logo(doc)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(titulo_curto)
    set_font(r1, size=28, bold=True, color=PRETO)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run(titulo)
    set_font(r2, size=18, bold=True, color=PRETO)


def add_meta_block(doc, content):
    """Bloco de metadados da reunião (empresa, data, hora, local, tipo) + objetivo."""
    empresa = content.get("empresa", "Singular")
    data = content.get("data", "")
    hora_inicio = content.get("hora_inicio", "")
    hora_fim = content.get("hora_fim", "")
    local = content.get("local", "")
    tipo = content.get("tipo_reuniao", "")
    objetivo = content.get("objetivo", "")

    # Linha 1: Empresa
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(f"Empresa: {empresa}")
    set_font(r1, size=11, color=CINZA)

    # Linha 2: Data | Horário | Local
    meta_bits = []
    if data:
        meta_bits.append(data)
    if hora_inicio:
        horario = hora_inicio + (f" – {hora_fim}" if hora_fim else "")
        meta_bits.append(horario)
    if local:
        meta_bits.append(local)
    if meta_bits:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run("  |  ".join(meta_bits))
        set_font(r2, size=11, color=CINZA)

    # Linha 3: Tipo de reunião
    if tipo:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(16)
        r3 = p3.add_run(f"Tipo: {tipo}")
        set_font(r3, size=11, color=CINZA)

    # Objetivo (se presente)
    if objetivo:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(18)
        r4a = p4.add_run("Objetivo: ")
        set_font(r4a, size=11, bold=True, color=PRETO)
        r4b = p4.add_run(objetivo)
        set_font(r4b, size=11, color=PRETO)


def add_section_heading(doc, titulo, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(titulo)
    set_font(r, size=size, bold=True, color=PRETO)


def add_bullet_list(doc, itens):
    for item in itens:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        run_b = p.add_run("•  ")
        set_font(run_b, size=11, bold=True, color=PRETO)
        run = p.add_run(item)
        set_font(run, size=11, color=PRETO)


def add_table_block(doc, colunas, linhas):
    total_cols = len(colunas)
    if total_cols == 0:
        return
    table = doc.add_table(rows=1 + len(linhas), cols=total_cols)
    table.autofit = True

    header = table.rows[0].cells
    for i, col_text in enumerate(colunas):
        cell = header[i]
        cell.text = ""
        set_cell_shading(cell, "1c1c1c")
        set_cell_borders(cell, color="d0d0d0")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(col_text)
        set_font(r, size=10, bold=True, color=BRANCO)

    for row_idx, row_data in enumerate(linhas, start=1):
        for col_idx in range(total_cols):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            set_cell_borders(cell, color="d0d0d0")
            text = row_data[col_idx] if col_idx < len(row_data) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            set_font(r, size=10, color=PRETO)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_participantes(doc, participantes):
    presentes = participantes.get("presentes", [])
    ausentes = participantes.get("ausentes", [])
    if not presentes and not ausentes:
        return
    add_section_heading(doc, "Participantes", size=14)
    if presentes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Presentes:")
        set_font(r, size=12, bold=True, color=PRETO)
        add_bullet_list(doc, presentes)
    if ausentes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Ausentes:")
        set_font(r, size=12, bold=True, color=PRETO)
        add_bullet_list(doc, ausentes)


def add_topico(doc, topico):
    add_section_heading(doc, topico["titulo"], size=13)

    for para in topico.get("discussao", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(para)
        set_font(r, size=11, color=PRETO)

    for tab in topico.get("tabelas", []):
        add_table_block(doc, tab["colunas"], tab["linhas"])

    decisoes = topico.get("decisoes", [])
    if decisoes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Decisões:")
        set_font(r, size=11, bold=True, color=PRETO)
        add_bullet_list(doc, decisoes)


def add_proxima_reuniao(doc, proxima):
    add_section_heading(doc, "Próxima Reunião", size=14)
    bits = []
    if proxima.get("data"):
        bits.append(proxima["data"])
    if proxima.get("hora"):
        bits.append(proxima["hora"])
    if proxima.get("local"):
        bits.append(proxima["local"])
    if bits:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("  |  ".join(bits))
        set_font(r, size=11, color=PRETO)
    pauta = proxima.get("pauta_prevista", [])
    if pauta:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Pauta prevista:")
        set_font(r, size=11, bold=True, color=PRETO)
        add_bullet_list(doc, pauta)


def add_assinaturas(doc, assinaturas):
    if not assinaturas:
        return
    add_section_heading(doc, "Assinaturas", size=14)
    for nome in assinaturas:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("_" * 45)
        set_font(r, size=11, color=CINZA_CLARO)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(nome)
        set_font(r2, size=11, color=PRETO)


def build(content, template_path, output_path):
    doc = Document(template_path)
    clear_document_body(doc)

    # Título
    add_title_block(doc, content.get("titulo_curto", "ATA"), content["titulo"])

    # Metadados
    add_meta_block(doc, content)

    # Participantes
    participantes = content.get("participantes")
    if participantes:
        add_participantes(doc, participantes)

    # Pauta
    pauta = content.get("pauta")
    if pauta:
        add_section_heading(doc, "Pauta", size=14)
        add_bullet_list(doc, pauta)

    # Tópicos (discussão + decisões)
    topicos = content.get("topicos", [])
    if topicos:
        add_section_heading(doc, "Tópicos Discutidos", size=14)
    for topico in topicos:
        add_topico(doc, topico)

    # Encaminhamentos (ação / responsável / prazo)
    enc = content.get("encaminhamentos")
    if enc and enc.get("linhas"):
        add_section_heading(doc, enc.get("titulo", "Encaminhamentos"), size=14)
        colunas = enc["linhas"][0]
        linhas = enc["linhas"][1:]
        add_table_block(doc, colunas, linhas)

    # Próxima reunião
    proxima = content.get("proxima_reuniao")
    if proxima:
        add_proxima_reuniao(doc, proxima)

    # Observações livres
    obs = content.get("observacoes")
    if obs:
        add_section_heading(doc, "Observações", size=14)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(obs)
        set_font(r, size=11, color=PRETO)

    # Assinaturas
    add_assinaturas(doc, content.get("assinaturas", []))

    doc.save(output_path)
    print(f"Gerado: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Uso: python build.py content.json output.docx")
        sys.exit(1)

    content_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    # Template oficial Singular (fonte Urbanist, header/footer com logo)
    template_path = Path(__file__).parent / "assets" / "template.docx"
    if not template_path.exists():
        print(f"Erro: template não encontrado em {template_path}")
        sys.exit(1)

    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    build(content, template_path, output_path)


if __name__ == "__main__":
    main()
