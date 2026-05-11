"""
Renderiza um template de contrato (markdown + Jinja2) para .docx no padrão Singular.

Uso CLI:
    python render_contrato.py --tipo nda-pf --vars dados.yml
    python render_contrato.py --tipo prestacao-servicos --vars cliente.yml --out /caminho/saida.docx

Tipos disponíveis: nda-pf, nda-pj, mou, prestacao-servicos, representacao-comercial

Estrutura do --vars (YAML):
    parte:
      nome: "Maria Silva"
      tipo: "PF"            # PF ou PJ
      documento: "123.456.789-00"
      rg: "1234567 SSP/DF"
      endereco: "Rua X, 100"
      bairro: "Asa Sul"
      cidade_uf: "Brasília/DF"
      cep: "70000-000"
      apelido: "MARIA"
    objeto: "consultoria estratégica em..."
    prazo_meses: 12
    data_assinatura: "26 de abril de 2026"
"""

import argparse
import datetime as dt
import re
import sys
from pathlib import Path

import yaml
from jinja2 import Environment, FileSystemLoader, StrictUndefined
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
TEMPLATES_DIR = SKILL_DIR / "templates"
PARTES_DIR = SKILL_DIR / "partes"

OUTPUT_BASE = Path(r"C:\Users\teste\Desktop\contratos")

VALID_TIPOS = {"nda-pf", "nda-pj", "mou", "prestacao-servicos", "representacao-comercial", "embaixador", "mutuo-royalty"}


def load_singular() -> dict:
    with open(PARTES_DIR / "singular.yml", encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_vars(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


def render_markdown(tipo: str, ctx: dict) -> str:
    env = Environment(
        loader=FileSystemLoader(TEMPLATES_DIR),
        undefined=StrictUndefined,
        trim_blocks=True,
        lstrip_blocks=True,
    )
    # Default filter pra evitar StrictUndefined em campos opcionais
    env.undefined = type("PermissiveUndefined", (), {
        "__bool__": lambda self: False,
        "__str__": lambda self: "",
    })
    # voltar ao default permissivo via ChainableUndefined
    from jinja2 import ChainableUndefined
    env.undefined = ChainableUndefined
    template = env.get_template(f"{tipo}.md")
    return template.render(**ctx)


# --- Markdown → docx (parser leve, suficiente pros nossos templates) ---

H1 = re.compile(r"^# (.+)$")
H2 = re.compile(r"^## (.+)$")
H3 = re.compile(r"^### (.+)$")
HR = re.compile(r"^---+$")
BULLET = re.compile(r"^- (.+)$")
BOLD = re.compile(r"\*\*([^*]+)\*\*")


def add_runs_with_bold(paragraph, text: str, base_size: int = 11):
    """Quebra o texto em runs preservando **bold**."""
    pos = 0
    for m in BOLD.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.size = Pt(base_size)
        run = paragraph.add_run(m.group(1))
        run.bold = True
        run.font.size = Pt(base_size)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(base_size)


def markdown_to_docx(md: str, out_path: Path, meta: dict):
    doc = Document()

    # Margens A4 estilo contrato
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # Style padrão
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if HR.match(line):
            # Separador visual — pula
            i += 1
            continue

        if H1.match(line):
            text = H1.match(line).group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            i += 1
            continue

        if H2.match(line):
            text = H2.match(line).group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue

        if H3.match(line):
            text = H3.match(line).group(1)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        if BULLET.match(line):
            text = BULLET.match(line).group(1)
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, text)
            i += 1
            continue

        # Parágrafo: junta linhas até encontrar linha vazia ou diretiva
        buf = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (not nxt.strip()) or HR.match(nxt) or H1.match(nxt) \
                    or H2.match(nxt) or H3.match(nxt) or BULLET.match(nxt):
                break
            buf.append(nxt)
            i += 1
        text = " ".join(buf)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs_with_bold(p, text)

    # Metadata no docx
    props = doc.core_properties
    props.author = "Singular Venture / Claude Master"
    props.title = meta.get("title", "Contrato Singular")
    props.subject = meta.get("tipo", "")
    props.comments = "Gerado por: Claude Master Desktop / Skill /contrato"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def normalize_apelido(parte: dict) -> dict:
    """Garante apelido em UPPER e fallbacks de campos."""
    p = dict(parte)
    if "apelido" not in p or not p["apelido"]:
        nome = p.get("nome", "PARTE")
        # Pega primeiro nome significativo
        first = nome.split(",")[0].split(" ")[0].upper()
        p["apelido"] = first
    p["apelido"] = p["apelido"].upper()
    return p


def slugify(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    return text.strip("-")[:50] or "contrato"


def main():
    ap = argparse.ArgumentParser(description="Renderiza contrato Singular em .docx")
    ap.add_argument("--tipo", required=True, choices=sorted(VALID_TIPOS))
    ap.add_argument("--vars", required=True, type=Path, help="YAML com variáveis")
    ap.add_argument("--out", type=Path, default=None, help="Caminho .docx (default: Desktop/contratos/...)")
    args = ap.parse_args()

    if not args.vars.exists():
        print(f"ERRO: arquivo de variáveis não encontrado: {args.vars}", file=sys.stderr)
        sys.exit(1)

    singular = load_singular()
    ctx = load_vars(args.vars) or {}
    ctx["singular"] = singular
    if "parte" in ctx:
        ctx["parte"] = normalize_apelido(ctx["parte"])
    if "data_assinatura" not in ctx:
        ctx["data_assinatura"] = dt.date.today().strftime("%d de %B de %Y")

    md = render_markdown(args.tipo, ctx)

    if args.out:
        out_path = args.out
    else:
        nome_curto = slugify(ctx.get("parte", {}).get("apelido", "contrato"))
        date_tag = dt.date.today().strftime("%Y-%m-%d")
        out_path = OUTPUT_BASE / args.tipo / f"{nome_curto}-{date_tag}.docx"

    markdown_to_docx(md, out_path, meta={"title": ctx.get("parte", {}).get("nome", ""), "tipo": args.tipo})

    print(f"OK | gerado: {out_path}")
    print(f"OK | dump_md: {out_path.with_suffix('.md')}")
    out_path.with_suffix(".md").write_text(md, encoding="utf-8")


if __name__ == "__main__":
    main()
