"""
build.py — Gera um Documento Formal Singular em .docx a partir de um JSON.

Uso:
    python build.py content.json output.docx

Reutiliza o template.docx oficial da Singular (em ~/.claude/skills/pop/template.docx)
com fontes Urbanist, header/footer da marca. Limpa o corpo e reconstrói com a estrutura
genérica de documento (metadados, TL;DR, seções com parágrafos/listas/tabelas/destaques,
conclusão, próximos passos, referências, assinaturas).

Schema esperado do content.json — ver SKILL.md.
"""
import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

LOGO_PATH = Path(__file__).parent / "assets" / "logo-singular.png"

URBANIST = "Urbanist"
PRETO = RGBColor(0x1C, 0x1C, 0x1C)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA = RGBColor(0x55, 0x55, 0x55)
CINZA_CLARO = RGBColor(0xD0, 0xD0, 0xD0)
CINZA_FUNDO = "F2F2F2"


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="d0d0d0", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def set_font(run, name=URBANIST, size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def clear_document_body(doc):
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def add_logo(doc):
    if not LOGO_PATH.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(1.8))


def add_title_block(doc, titulo_curto, titulo, subtitulo=None):
    add_logo(doc)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(titulo_curto)
    set_font(r1, size=24, bold=True, color=PRETO)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4 if subtitulo else 16)
    r2 = p2.add_run(titulo)
    set_font(r2, size=18, bold=True, color=PRETO)

    if subtitulo:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(16)
        r3 = p3.add_run(subtitulo)
        set_font(r3, size=12, italic=True, color=CINZA)


def add_meta_block(doc, content):
    """Bloco de metadados (empresa | autor | destinatário | data)."""
    bits = []
    if content.get("empresa"):
        bits.append(f"Empresa: {content['empresa']}")
    if content.get("autor"):
        bits.append(f"Autor: {content['autor']}")
    if content.get("destinatario"):
        bits.append(f"Para: {content['destinatario']}")
    if content.get("data"):
        bits.append(content["data"])

    if bits:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run("  |  ".join(bits))
        set_font(r, size=10, color=CINZA)


def add_horizontal_rule(doc, space_after=10):
    """Linha cinza fina para separar blocos."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "D0D0D0")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_tldr_block(doc, tldr_text):
    """TL;DR em caixa cinza claro logo após o título."""
    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_after = Pt(4)
    r_label = label_p.add_run("TL;DR")
    set_font(r_label, size=10, bold=True, color=CINZA)

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, CINZA_FUNDO)
    set_cell_borders(cell, color="e6e6e6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(tldr_text)
    set_font(r, size=11, color=PRETO)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)


def add_section_heading(doc, titulo, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    space_before = {1: 14, 2: 10, 3: 6}
    space_after = {1: 8, 2: 6, 3: 4}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before.get(level, 6))
    p.paragraph_format.space_after = Pt(space_after.get(level, 4))
    r = p.add_run(titulo)
    set_font(r, size=size, bold=True, color=PRETO)


def add_paragraph_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=11, color=PRETO)


def add_list_block(doc, lista):
    tipo = lista.get("tipo", "bullet")
    itens = lista.get("itens", [])
    for idx, item in enumerate(itens, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        marker = f"{idx}.  " if tipo == "numbered" else "•  "
        run_b = p.add_run(marker)
        set_font(run_b, size=11, bold=(tipo == "numbered"), color=PRETO)
        run = p.add_run(item)
        set_font(run, size=11, color=PRETO)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_table_block(doc, colunas, linhas):
    total_cols = len(colunas)
    if total_cols == 0:
        return
    table = doc.add_table(rows=1 + len(linhas), cols=total_cols)
    table.autofit = True

    header = table.rows[0].cells
    for i, col_text in enumerate(colunas):
        cell = header[i]
        cell.text = ""
        set_cell_shading(cell, "1c1c1c")
        set_cell_borders(cell, color="d0d0d0")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(col_text)
        set_font(r, size=10, bold=True, color=BRANCO)

    for row_idx, row_data in enumerate(linhas, start=1):
        for col_idx in range(total_cols):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            set_cell_borders(cell, color="d0d0d0")
            text = row_data[col_idx] if col_idx < len(row_data) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            set_font(r, size=10, color=PRETO)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_destaque_block(doc, texto):
    """Callout em caixa cinza claro, fonte itálica."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, CINZA_FUNDO)
    set_cell_borders(cell, color="d0d0d0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texto)
    set_font(r, size=11, italic=True, color=PRETO)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def render_secao(doc, secao, level=1):
    titulo = secao.get("titulo", "")
    if titulo:
        add_section_heading(doc, titulo, level=level)

    for paragrafo in secao.get("paragrafos", []):
        add_paragraph_text(doc, paragrafo)

    for lista in secao.get("listas", []):
        add_list_block(doc, lista)

    for tab in secao.get("tabelas", []):
        add_table_block(doc, tab.get("colunas", []), tab.get("linhas", []))

    if secao.get("destaque"):
        add_destaque_block(doc, secao["destaque"])

    for sub in secao.get("subsecoes", []):
        render_secao(doc, sub, level=min(level + 1, 3))


def add_conclusao(doc, texto):
    add_section_heading(doc, "Conclusão", level=1)
    add_paragraph_text(doc, texto)


def add_proximos_passos(doc, passos):
    add_section_heading(doc, "Próximos Passos", level=1)
    add_list_block(doc, {"tipo": "numbered", "itens": passos})


def add_referencias(doc, refs):
    add_section_heading(doc, "Referências", level=1)
    add_list_block(doc, {"tipo": "bullet", "itens": refs})


def add_autor_rodape(doc, autor):
    """Rodapé simples no fim do documento, apenas com autor — sem linha de assinatura."""
    if not autor:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(autor)
    set_font(r, size=10, italic=True, color=CINZA)


def build(content, template_path, output_path):
    doc = Document(template_path)
    clear_document_body(doc)

    add_title_block(
        doc,
        content.get("titulo_curto", "DOCUMENTO"),
        content["titulo"],
        content.get("subtitulo"),
    )

    add_meta_block(doc, content)
    add_horizontal_rule(doc, space_after=14)

    if content.get("tldr"):
        add_tldr_block(doc, content["tldr"])

    for secao in content.get("secoes", []):
        render_secao(doc, secao, level=1)

    if content.get("conclusao"):
        add_conclusao(doc, content["conclusao"])

    if content.get("proximos_passos"):
        add_proximos_passos(doc, content["proximos_passos"])

    if content.get("referencias"):
        add_referencias(doc, content["referencias"])

    if content.get("autor_rodape"):
        add_autor_rodape(doc, content["autor_rodape"])

    doc.save(output_path)
    print(f"Gerado: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Uso: python build.py content.json output.docx")
        sys.exit(1)

    content_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    template_path = Path(__file__).parent.parent / "pop" / "template.docx"
    if not template_path.exists():
        print(f"Erro: template não encontrado em {template_path}")
        sys.exit(1)

    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    if "titulo" not in content:
        print("Erro: campo obrigatório 'titulo' ausente no content.json")
        sys.exit(1)
    if "secoes" not in content or not content["secoes"]:
        print("Erro: 'secoes' deve conter ao menos uma seção")
        sys.exit(1)

    build(content, template_path, output_path)


if __name__ == "__main__":
    main()
