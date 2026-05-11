"""
build.py — Gera um POP Singular em .docx a partir de um JSON de conteudo.

Uso:
    python build.py content.json output.docx

O script abre template.docx (reference da Singular com fonts/header/footer/estilos
já configurados em Urbanist + paleta preta), LIMPA o corpo do documento, e
reconstrói com o conteúdo fornecido. Header e footer são preservados
automaticamente porque são parte da seção.

Estrutura esperada do content.json:

{
  "titulo_curto": "POP",
  "titulo": "Processo Operacional de [X]",
  "empresa": "Singular",
  "versao": "1.0",
  "data": "15 de abril de 2026",
  "objetivo": "Frase objetivo do processo.",
  "perfil_secao": {
    "titulo": "Perfil do Cliente",
    "bullets": ["item 1", "item 2", ...]
  },
  "fluxo_titulo": "Fluxo em N Passos",
  "passos": [
    {
      "titulo": "Passo 1 — Nome",
      "paragrafos": ["texto de introducao", ...],
      "tabelas": [
        {
          "colunas": ["Cenario", "Resposta padrao"],
          "linhas": [["a", "b"], ["c", "d"]]
        }
      ],
      "listas": [
        {"titulo": "Regras:", "itens": ["item", "item"]}
      ],
      "acao_final": "Frase destacada ao fim do passo (opcional)"
    }
  ],
  "checklist": {
    "titulo": "Checklist do Vendedor",
    "secoes": [
      {"titulo": "Antes da visita", "itens": ["item", "item"]},
      {"titulo": "Após a visita", "itens": ["item"]}
    ]
  },
  "contato": {
    "titulo": "Contato Comercial",
    "linhas": [["Canal", "Contato"], ["WhatsApp", "61 9 9126-1177"]]
  }
}
"""
import json
import sys
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

URBANIST = "Urbanist"
PRETO = RGBColor(0x1C, 0x1C, 0x1C)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA = RGBColor(0x55, 0x55, 0x55)
CINZA_CLARO = RGBColor(0xD0, 0xD0, 0xD0)


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="d0d0d0", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def set_font(run, name=URBANIST, size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # Garante que east asia/cs também usem a fonte
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def clear_document_body(doc):
    """Remove todos paragrafos/tabelas do corpo, preservando section properties."""
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def add_paragraph(doc, text, size=11, bold=False, align=None, color=None, space_after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color or PRETO)
    return p


def add_title_block(doc, titulo_curto, titulo):
    """Cabecalho visual: POP em cobre/pequeno + título grande."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(titulo_curto)
    set_font(r1, size=28, bold=True, color=PRETO)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run(titulo)
    set_font(r2, size=18, bold=True, color=PRETO)


def add_meta_block(doc, empresa, versao, data, objetivo):
    """Bloco de metadados centralizado em cinza, seguido do objetivo."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(f"Empresa: {empresa}")
    set_font(r1, size=11, color=CINZA)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run(f"Versão {versao}  |  {data}")
    set_font(r2, size=11, color=CINZA)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(18)
    r3a = p3.add_run("Objetivo: ")
    set_font(r3a, size=11, bold=True, color=PRETO)
    r3b = p3.add_run(objetivo)
    set_font(r3b, size=11, color=PRETO)


def add_section_heading(doc, titulo, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(titulo)
    set_font(r, size=size, bold=True, color=PRETO)


def add_bullet_list(doc, itens):
    for item in itens:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        run_b = p.add_run("•  ")
        set_font(run_b, size=11, bold=True, color=PRETO)
        run = p.add_run(item)
        set_font(run, size=11, color=PRETO)


def add_table_block(doc, colunas, linhas):
    """Tabela com cabecalho preto (fill 1c1c1c) + texto branco."""
    total_cols = len(colunas)
    if total_cols == 0:
        return
    table = doc.add_table(rows=1 + len(linhas), cols=total_cols)
    table.autofit = True

    # Cabeçalho
    header = table.rows[0].cells
    for i, col_text in enumerate(colunas):
        cell = header[i]
        cell.text = ""
        set_cell_shading(cell, "1c1c1c")
        set_cell_borders(cell, color="d0d0d0")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(col_text)
        set_font(r, size=10, bold=True, color=BRANCO)

    # Corpo
    for row_idx, row_data in enumerate(linhas, start=1):
        for col_idx in range(total_cols):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            set_cell_borders(cell, color="d0d0d0")
            text = row_data[col_idx] if col_idx < len(row_data) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            set_font(r, size=10, color=PRETO)

    # Espaço após tabela
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_passo(doc, passo):
    add_section_heading(doc, passo["titulo"], size=13)
    for para in passo.get("paragrafos", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(para)
        set_font(r, size=11, color=PRETO)
    for tab in passo.get("tabelas", []):
        add_table_block(doc, tab["colunas"], tab["linhas"])
    for lst in passo.get("listas", []):
        if lst.get("titulo"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(lst["titulo"])
            set_font(r, size=11, bold=True, color=PRETO)
        add_bullet_list(doc, lst["itens"])
    if passo.get("acao_final"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        ra = p.add_run("Ação: ")
        set_font(ra, size=11, bold=True, color=PRETO)
        rb = p.add_run(passo["acao_final"])
        set_font(rb, size=11, color=PRETO)


def build(content, template_path, output_path):
    doc = Document(template_path)
    clear_document_body(doc)

    # Title block
    add_title_block(doc, content.get("titulo_curto", "POP"), content["titulo"])

    # Metadata + objetivo
    add_meta_block(
        doc,
        content.get("empresa", "Singular"),
        content.get("versao", "1.0"),
        content.get("data", ""),
        content["objetivo"],
    )

    # Perfil (ou seção similar introdutória)
    perfil = content.get("perfil_secao")
    if perfil:
        add_section_heading(doc, perfil["titulo"], size=14)
        add_bullet_list(doc, perfil.get("bullets", []))

    # Fluxo
    if content.get("fluxo_titulo"):
        add_section_heading(doc, content["fluxo_titulo"], size=14)
    for passo in content.get("passos", []):
        add_passo(doc, passo)

    # Checklist
    cl = content.get("checklist")
    if cl:
        add_section_heading(doc, cl["titulo"], size=14)
        for secao in cl.get("secoes", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(secao["titulo"])
            set_font(r, size=12, bold=True, color=PRETO)
            add_bullet_list(doc, secao.get("itens", []))

    # Contato
    ct = content.get("contato")
    if ct:
        add_section_heading(doc, ct["titulo"], size=14)
        if ct.get("linhas"):
            colunas = ct["linhas"][0]
            linhas = ct["linhas"][1:]
            add_table_block(doc, colunas, linhas)

    doc.save(output_path)
    print(f"Gerado: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Uso: python build.py content.json output.docx")
        sys.exit(1)

    content_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    template_path = Path(__file__).parent / "template.docx"

    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    build(content, template_path, output_path)


if __name__ == "__main__":
    main()
